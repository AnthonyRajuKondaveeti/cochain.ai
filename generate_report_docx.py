import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create a w:r element and a new w:rPr element
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run ()
    r._r.append (hyperlink)

    # A workaround for the color of the hyperlink
    r.font.color.rgb = RGBColor(0, 0, 255)
    r.font.underline = True

    return hyperlink

def process_inline_formatting(paragraph, text):
    # Handle bold **text**
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)

def create_word_doc(md_file_path, output_path):
    doc = Document()
    
    # Set styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Title
    doc.add_heading('Reinforcement Learning Implementation and Training Report', 0)
    
    with open(md_file_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    in_code_block = False
    code_content = []
    in_table = False
    table_lines = []
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        original_line = lines[i] # Keep original for indentation check if needed
        
        # Skip empty lines if not in code block
        if not line and not in_code_block:
            i += 1
            continue
            
        # 1. Handle Code Blocks
        if line.startswith('```'):
            if in_code_block:
                # End of code block
                in_code_block = False
                # Add code paragraph
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                run = p.add_run('\n'.join(code_content))
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(50, 50, 50)
                code_content = []
            else:
                # Start of code block
                in_code_block = True
            i += 1
            continue
            
        if in_code_block:
            code_content.append(original_line.rstrip())
            i += 1
            continue

        # 2. Handle Tables
        if line.startswith('|'):
            if not in_table:
                # Check if it's a valid table start (next line should be separator)
                if i + 1 < len(lines) and '---' in lines[i+1]:
                    in_table = True
                    table_lines = [line]
                else:
                    # Just text starting with |
                    p = doc.add_paragraph(line)
            else:
                table_lines.append(line)
            i += 1
            continue
        
        if in_table:
            # End of table detected (line doesn't start with |)
            in_table = False
            # Process table
            # Filter out separator lines
            data_rows = [l for l in table_lines if '---' not in l]
            if data_rows:
                # Determine columns
                header = [c.strip() for c in data_rows[0].strip('|').split('|')]
                table = doc.add_table(rows=len(data_rows), cols=len(header))
                table.style = 'Table Grid'
                
                # Fill data
                for r_idx, row_text in enumerate(data_rows):
                    cells = [c.strip() for c in row_text.strip('|').split('|')]
                    row_cells = table.rows[r_idx].cells
                    for c_idx, cell_text in enumerate(cells):
                        if c_idx < len(row_cells):
                            row_cells[c_idx].text = cell_text
                            # Bold headers
                            if r_idx == 0:
                                for paragraph in row_cells[c_idx].paragraphs:
                                    for run in paragraph.runs:
                                        run.font.bold = True
            table_lines = []
            # Don't skip the current line, process it as normal text
            # But wait, we already incremented i inside the loop? No, we are here because line doesn't start with |
            # So we fall through to normal processing
        
        # 3. Handle Headers
        if line.startswith('#'):
            level = len(line.split(' ')[0])
            text = line.lstrip('#').strip()
            # Docx supports heading levels 1-9
            # Map level 1 (#) to Heading 1, etc.
            # But our doc title is level 0 (Title), so # -> Heading 1
            if level > 9: level = 9
            doc.add_heading(text, level=level)
            i += 1
            continue
            
        # 4. Handle Lists
        if line.startswith('- ') or line.startswith('* '):
            text = line[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            process_inline_formatting(p, text)
            i += 1
            continue
            
        if re.match(r'^\d+\.', line):
            # Numbered list
            parts = line.split('.', 1)
            if len(parts) > 1:
                text = parts[1].strip()
                p = doc.add_paragraph(style='List Number')
                process_inline_formatting(p, text)
                i += 1
                continue

        # 5. Normal Text
        p = doc.add_paragraph()
        process_inline_formatting(p, line)
        i += 1

    doc.save(output_path)
    print(f"Successfully generated {output_path}")

if __name__ == "__main__":
    create_word_doc(
        r'e:\5MDS\Project\project\RL_TRAINING_REPORT.md',
        r'e:\5MDS\Project\project\RL_TRAINING_REPORT.docx'
    )
